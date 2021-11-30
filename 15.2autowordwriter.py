import docx
from docx.shared import RGBColor
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import time

file = docx.Document()
file.styles['Normal'].font.name = u'微软雅黑' #可换成word里面任意字体
file.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

# --- 创建一个封面 ---
p = file.add_paragraph() # 创建一个段落
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # 居中设置
p.paragraph_format.space_before = Pt(200) # 段前距为200，这个是测试出来的
p.paragraph_format.space_after = Pt(30) # 段后距为30，这个也是测试出来的
run = p.add_run('华小智舆情报告') # 在段落里添加内容
font = run.font # 设置字体
font.color.rgb = RGBColor(54, 95, 145) # 颜色设置，这里是用RGB颜色
font.size = Pt(42) # 字体大小设置，和word里面的字号相对应

# 在封面上添加日期
year = time.strftime("%Y")
month = time.strftime("%m")
day = time.strftime("%d")
today = year + '年' + month + '月' + day + '日' # 构造当天日期
p = file.add_paragraph() # 新建一个段落
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(today) # 在段落中输入当天日期
font = run.font
font.color.rgb = RGBColor(54, 95, 145)
font.size = Pt(26)

# --- 开始写正文 ---
file.add_page_break()  # 添加分页符
p = file.add_paragraph()  # 设置正文标题
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
run = p.add_run('第一部分 阿里巴巴舆情报告')
run.font.color.rgb = RGBColor(54, 95, 145)  # 字体颜色设置
run.font.size = Pt(22)  # 字体大小设置

# 把Word文档保存，注意需提前创建好保存文件夹
file.save('/Users/katherinewang/Desktop/dltest/华小智舆情报告.docx')
print('华小智舆情报告生成完毕')