# === 用python控制doc ===
import docx
from docx import Document
from docx.shared import Inches

document = Document() # NOTE: initialize

# --- heading & paragraph ---
document.add_heading('Document Title', 0)  # 添加标题，但是这个默认会有下划线

# --- 设置中文字体 ---
from docx.oxml.ns import qn
document.styles['Normal'].font.name = u'微软雅黑'  # 可换成word里面任意字体
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 这边也得填一下字体名称

document.add_paragraph('漫天的我落在枫叶上雪花上')  # 这个方式也可以作为标题
pp = document.add_paragraph()
run = pp.add_run('我喜欢你')
font = run.font  # NOTE：不可以直接pp.font，单纯的段落是没有字体属性的

# --- 设置字体大小和颜色 ---
from docx.shared import Pt
font.size = Pt(26)
from docx.shared import RGBColor
font.color.rgb = RGBColor(54,95,145)

# --- 设置粗体斜体下划线 ---
p = document.add_paragraph()
run = p.add_run('漫天的我落在枫叶上雪花上')
font = run.font
font.bold = True #粗体
font.italic = True #斜体
font.underline = True #下划线
## 或者这么写
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True  # 如果想在后面添加，并不想新建段落时，使用add.run
p.add_run(' and some ')
p.add_run('italic.').italic = True

# --- 文中段落设置 ---
from docx.enum.text import WD_ALIGN_PARAGRAPH
p = document.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # NOTE：居中
p.add_run('而你在想我')

from docx.shared import Inches
p = document.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.32)
p.add_run('设置首行缩进示例文字')  # NOTE：首行缩进

from docx.shared import Pt
p = document.add_paragraph()
p.paragraph_format.line_spacing = Pt(16)  # 行间距，16磅表示三号字体对应的宽度
p.add_run('设置行间距示例文字')  # NOTE：首行缩进

from docx.shared import Pt
p = document.add_paragraph()
p.paragraph_format.space_before = Pt(14)  # 段前距,14磅表示4号字体
p.paragraph_format.space_after = Pt(14)  # 段后距
p.add_run('设置段前段后距示例文字')  # NOTE：设置段前距和段后距

# NOTE：设置引用和段落序号
document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph('first item in unordered list', style='List Bullet')
document.add_paragraph('second item in unordered list', style='List Bullet')

document.add_paragraph('first item in ordered list', style='List Number')
document.add_paragraph('second item in ordered list', style='List Number')

document.add_page_break()  # NOTE: add a page

# --- table set up ---
## method 1
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

## method 2
table = document.add_table(rows=2, cols=3, style='Light Shading Accent 1')
table.cell(0, 0).text = '第一句'  # 第一行第一列
table.cell(0, 1).text = '第二句'  # 第一行第二列
table.cell(0, 2).text = '第三句'  # 第一行第三列
table.cell(1, 0).text = '克制'  # 第二行第一列
table.cell(1, 1).text = '再克制'  # 第二行第二列
table.cell(1, 2).text = '"在吗"'  # 第三行第三列

# --- 图片设置 ---
document.add_picture('/Users/katherinewang/Desktop/dltest/up.jpg', width=Inches(1.25), height=Inches(1.25))
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


document.save('/Users/katherinewang/Desktop/dltest/templet.docx')
print('Word生成完毕')

# === use python to read from doc ===
readfile = docx.Document('/Users/katherinewang/Desktop/dltest/templet.docx') # 打开文件
content = [] #
for paragraph in readfile.paragraphs:
    print(paragraph.text)  # 打印各段落内容文本，注意不要忘写.text
    content.append(paragraph.text)
content = ' '.join(content) #其中单引号里的内容表示利用空格进行连接
print(content)

